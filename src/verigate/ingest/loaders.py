"""File loaders — turn a customer document into plain text for extraction.

Each loader is deterministic and offline. A file that cannot be read or
parsed raises :class:`LoaderError` with a precise message (file + cause),
chaining the original exception — never swallowed, never `errors='replace'`
(a silently mangled corpus would "verify" garbage). PDF and DOCX support
live behind the optional ``verigate[ingest]`` extra (D-009): the imports are
lazy and a missing dependency raises a LoaderError telling the user what to
install.
"""

from __future__ import annotations

import csv
import io
import json
from pathlib import Path

#: Extensions :func:`load_file` can dispatch on (lowercase, dot included).
SUPPORTED_EXTENSIONS = frozenset({".txt", ".md", ".csv", ".json", ".pdf", ".docx"})


class LoaderError(Exception):
    """A document could not be loaded (unreadable, corrupt, unsupported)."""


def load_file(path: Path) -> str:
    """Load `path` as plain text, dispatching on the (case-insensitive)
    suffix. Raises :class:`LoaderError` on unreadable/corrupt input or an
    unsupported extension."""
    suffix = path.suffix.lower()
    if suffix in (".txt", ".md"):
        return _load_text(path)
    if suffix == ".csv":
        return _load_csv(path)
    if suffix == ".json":
        return _load_json(path)
    if suffix == ".pdf":
        return _load_pdf(path)
    if suffix == ".docx":
        return _load_docx(path)
    raise LoaderError(f"{path}: unsupported extension {suffix!r}")


def _load_text(path: Path) -> str:
    """Strict UTF-8 read. A corrupt encoding is an error, not a guess."""
    try:
        return path.read_text(encoding="utf-8")
    except UnicodeDecodeError as exc:
        raise LoaderError(f"{path}: not valid UTF-8 ({exc})") from exc
    except OSError as exc:
        raise LoaderError(f"{path}: unreadable ({exc})") from exc


def _load_csv(path: Path) -> str:
    """Header-aware rendering: one line per data row, cells rendered as
    ``col: value | col: value``. Values are kept verbatim (no numeric
    parsing) so the number/reference extractors see exactly the source
    text. Cells beyond the header width are labeled ``colN``."""
    text = _load_text(path)
    try:
        rows = list(csv.reader(io.StringIO(text)))
    except csv.Error as exc:
        raise LoaderError(f"{path}: malformed CSV ({exc})") from exc
    if not rows:
        return ""
    header = rows[0]
    lines: list[str] = []
    for row in rows[1:]:
        parts = []
        for i, cell in enumerate(row):
            label = header[i] if i < len(header) else f"col{i + 1}"
            parts.append(f"{label}: {cell}")
        lines.append(" | ".join(parts))
    return "\n".join(lines)


def _load_json(path: Path) -> str:
    """Flatten the JSON document to ``dotted.path: value`` lines, list items
    as ``path[i]``. String values are rendered raw (no quotes) so embedded
    references extract cleanly; other scalars via ``json.dumps``. Non-dict
    roots are handled (a list root yields ``[i]: …`` lines, a scalar root a
    single value line)."""
    text = _load_text(path)
    try:
        data = json.loads(text)
    except json.JSONDecodeError as exc:
        raise LoaderError(f"{path}: invalid JSON ({exc})") from exc
    return "\n".join(_flatten_json(data, ""))


def _flatten_json(value: object, prefix: str) -> list[str]:
    """Recursive flattener; iteration follows document order, which is
    deterministic for identical input bytes."""
    if isinstance(value, dict):
        lines: list[str] = []
        for key, sub in value.items():
            sub_prefix = f"{prefix}.{key}" if prefix else str(key)
            lines.extend(_flatten_json(sub, sub_prefix))
        return lines
    if isinstance(value, list):
        lines = []
        for i, sub in enumerate(value):
            lines.extend(_flatten_json(sub, f"{prefix}[{i}]"))
        return lines
    rendered = value if isinstance(value, str) else json.dumps(value)
    return [f"{prefix}: {rendered}" if prefix else f"{rendered}"]


def _load_pdf(path: Path) -> str:
    """Extract the text of every page, joined by newlines (lazy pypdf)."""
    try:
        from pypdf import PdfReader
        from pypdf.errors import PyPdfError
    except ImportError as exc:
        raise LoaderError(
            f"{path}: PDF support requires the optional dependency pypdf — "
            "install verigate[ingest]"
        ) from exc
    try:
        reader = PdfReader(str(path))
        return "\n".join(page.extract_text() for page in reader.pages)
    except PyPdfError as exc:
        raise LoaderError(f"{path}: corrupt or unreadable PDF ({exc})") from exc
    except OSError as exc:
        raise LoaderError(f"{path}: unreadable ({exc})") from exc


def _load_docx(path: Path) -> str:
    """Paragraph texts plus table cell texts, joined by newlines (lazy
    python-docx)."""
    try:
        import docx
        from docx.opc.exceptions import PackageNotFoundError
    except ImportError as exc:
        raise LoaderError(
            f"{path}: DOCX support requires the optional dependency python-docx — "
            "install verigate[ingest]"
        ) from exc
    try:
        document = docx.Document(str(path))
    except PackageNotFoundError as exc:
        raise LoaderError(f"{path}: corrupt or unreadable DOCX ({exc})") from exc
    except OSError as exc:
        raise LoaderError(f"{path}: unreadable ({exc})") from exc
    parts: list[str] = [paragraph.text for paragraph in document.paragraphs]
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                parts.append(cell.text)
    return "\n".join(parts)
