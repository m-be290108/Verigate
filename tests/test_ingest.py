"""Tests for the ingestion pipeline (loaders + ingest_folder).

Everything here is deterministic and offline: real files in tmp_path,
real sqlite, no network, no LLM.
"""

from __future__ import annotations

from pathlib import Path

import pytest

from verigate.canonical import canonical_entity, canonical_ref, canonical_text
from verigate.corpus import CorpusDB
from verigate.ingest.ingestor import IngestResult, ingest_folder
from verigate.ingest.loaders import SUPPORTED_EXTENSIONS, LoaderError, load_file

# --------------------------------------------------------------- helpers


def _ingest(folder: Path, tmp_path: Path, **kwargs) -> tuple[IngestResult, Path]:
    db_path = tmp_path / "corpus.db"
    return ingest_folder(folder, db_path, **kwargs), db_path


def _minimal_pdf(text: str) -> bytes:
    """Handcraft a one-page PDF with a single text object (ASCII only,
    no parentheses/backslashes in `text`)."""
    stream = f"BT /F1 12 Tf 72 720 Td ({text}) Tj ET".encode("ascii")
    objects = [
        b"<< /Type /Catalog /Pages 2 0 R >>",
        b"<< /Type /Pages /Kids [3 0 R] /Count 1 >>",
        (
            b"<< /Type /Page /Parent 2 0 R /MediaBox [0 0 612 792] "
            b"/Resources << /Font << /F1 5 0 R >> >> /Contents 4 0 R >>"
        ),
        b"<< /Length "
        + str(len(stream)).encode("ascii")
        + b" >>\nstream\n"
        + stream
        + b"\nendstream",
        b"<< /Type /Font /Subtype /Type1 /BaseFont /Helvetica >>",
    ]
    out = bytearray(b"%PDF-1.4\n")
    offsets = []
    for i, obj in enumerate(objects, start=1):
        offsets.append(len(out))
        out += f"{i} 0 obj\n".encode("ascii") + obj + b"\nendobj\n"
    xref_pos = len(out)
    out += b"xref\n0 " + str(len(objects) + 1).encode("ascii") + b"\n0000000000 65535 f \n"
    for off in offsets:
        out += f"{off:010d} 00000 n \n".encode("ascii")
    out += (
        b"trailer\n<< /Size "
        + str(len(objects) + 1).encode("ascii")
        + b" /Root 1 0 R >>\nstartxref\n"
        + str(xref_pos).encode("ascii")
        + b"\n%%EOF\n"
    )
    return bytes(out)


# ------------------------------------------------- end-to-end (sample dir)


def test_e2e_sample_doc_count_fingerprint_no_skips(sample_corpus_dir, tmp_path):
    result, _ = _ingest(sample_corpus_dir, tmp_path)
    assert result.n_docs == 3
    assert result.fingerprint  # non-empty
    assert result.skipped == ()


def test_e2e_sample_refs_include_sku_set(sample_corpus_dir, tmp_path):
    _, db_path = _ingest(sample_corpus_dir, tmp_path)
    with CorpusDB(db_path) as db:
        assert db.has_reference(canonical_ref("AP-3000-X")) is not None
        assert db.has_reference(canonical_ref("HF-MINI-2")) is not None


def test_e2e_sample_numbers_include_money(sample_corpus_dir, tmp_path):
    _, db_path = _ingest(sample_corpus_dir, tmp_path)
    with CorpusDB(db_path) as db:
        assert db.has_number("money:EUR:249.99", kind="money") is not None
        assert db.has_number("money:EUR:39.5", kind="money") is not None


def test_e2e_sample_entities_from_md_heading_and_csv(sample_corpus_dir, tmp_path):
    result, db_path = _ingest(sample_corpus_dir, tmp_path)
    with CorpusDB(db_path) as db:
        # md heading minus parenthesized '(SKU AP-3000-X)' part:
        assert db.has_entity(canonical_entity("AquaPump 3000")) == "catalog.md"
        assert db.has_entity(canonical_entity("HydroFilter Mini")) == "catalog.md"
        canonicals = [c for c, _raw in db.entities()]
        assert canonicals == ["aquapump 3000", "hydrofilter mini"]
        # 2 heading rows (catalog.md) + 2 name-column rows (products.csv):
        assert result.n_entities == 4
        # CSV name-column provenance: drop catalog.md, the CSV rows remain.
        db.delete_document("catalog.md")
        assert db.has_entity(canonical_entity("AquaPump 3000")) == "products.csv"


def test_e2e_sample_warranty_quote_findable(sample_corpus_dir, tmp_path):
    _, db_path = _ingest(sample_corpus_dir, tmp_path)
    quote = "This product is covered for 24 months from the date of purchase."
    with CorpusDB(db_path) as db:
        assert db.contains_text(canonical_text(quote)) == "catalog.md"


def test_e2e_sample_verify_corpus_clean(sample_corpus_dir, tmp_path):
    _, db_path = _ingest(sample_corpus_dir, tmp_path)
    with CorpusDB(db_path) as db:
        assert db.verify_corpus() == (True, [])


def test_e2e_runs_fully_offline(sample_corpus_dir, tmp_path, no_network):
    result, _ = _ingest(sample_corpus_dir, tmp_path)
    assert result.fingerprint


# ------------------------------------------------------------ determinism


def test_determinism_two_fresh_dbs_identical_fingerprints(sample_corpus_dir, tmp_path):
    r1 = ingest_folder(sample_corpus_dir, tmp_path / "a.db")
    r2 = ingest_folder(sample_corpus_dir, tmp_path / "b.db")
    assert r1.fingerprint == r2.fingerprint
    assert r1 == r2


def test_reingest_same_db_idempotent(sample_corpus_dir, tmp_path):
    db_path = tmp_path / "corpus.db"
    r1 = ingest_folder(sample_corpus_dir, db_path)
    r2 = ingest_folder(sample_corpus_dir, db_path)
    assert r1.fingerprint == r2.fingerprint
    assert (r1.n_docs, r1.n_refs, r1.n_numbers, r1.n_entities) == (
        r2.n_docs,
        r2.n_refs,
        r2.n_numbers,
        r2.n_entities,
    )
    with CorpusDB(db_path) as db:
        assert db.verify_corpus() == (True, [])


# ------------------------------------------------------- pruned (explicit)


def _two_file_folder(tmp_path: Path) -> Path:
    folder = tmp_path / "src"
    folder.mkdir()
    (folder / "keep.md").write_text("# Keep\nWarranty: 24 months.\n", encoding="utf-8")
    (folder / "obsolete.md").write_text(
        "# Obsolete\nOld price: 999.99 EUR. SKU OB-1234-Z.\n", encoding="utf-8"
    )
    return folder


def test_reingest_prunes_deleted_file_and_reports_it(tmp_path):
    """Regression (2026-06-10 review, D-011): a file removed from the
    trusted folder must leave the corpus at the next ingest — doc gone, FTS
    finds nothing from it, fingerprint changes, prune is reported."""
    folder = _two_file_folder(tmp_path)
    db_path = tmp_path / "corpus.db"
    r1 = ingest_folder(folder, db_path)
    assert r1.n_docs == 2
    assert r1.pruned == ()

    (folder / "obsolete.md").unlink()
    r2 = ingest_folder(folder, db_path)

    assert r2.pruned == ("obsolete.md",)
    assert r2.n_docs == 1
    assert r2.fingerprint != r1.fingerprint
    with CorpusDB(db_path) as db:
        assert db.doc_ids() == ["keep.md"]
        assert db.search("OB-1234-Z") == []
        assert db.search("Obsolete") == []
        assert db.has_reference(canonical_ref("OB-1234-Z")) is None
        assert db.has_number("money:EUR:999.99", kind="money") is None
        assert db.contains_text(canonical_text("Old price: 999.99 EUR")) is None
        assert db.verify_corpus() == (True, [])


def test_update_ingest_matches_fresh_build_after_deletion(tmp_path):
    """The fingerprint in a report must be reproducible from the source
    folder alone: updating an existing db equals a fresh build of the
    identical folder (D-011)."""
    folder = _two_file_folder(tmp_path)
    db_path = tmp_path / "updated.db"
    ingest_folder(folder, db_path)
    (folder / "obsolete.md").unlink()
    updated = ingest_folder(folder, db_path)
    fresh = ingest_folder(folder, tmp_path / "fresh.db")
    assert updated.fingerprint == fresh.fingerprint
    assert (updated.n_docs, updated.n_refs, updated.n_numbers, updated.n_entities) == (
        fresh.n_docs,
        fresh.n_refs,
        fresh.n_numbers,
        fresh.n_entities,
    )


def test_glossary_doc_counts_as_seen_never_pruned(tmp_path):
    folder = tmp_path / "src"
    folder.mkdir()
    (folder / "note.txt").write_text("nothing here\n", encoding="utf-8")
    (folder / "glossary.yaml").write_text("- AquaPump 3000\n", encoding="utf-8")
    db_path = tmp_path / "corpus.db"
    r1 = ingest_folder(folder, db_path)
    r2 = ingest_folder(folder, db_path)
    assert r2.pruned == ()
    assert r2.fingerprint == r1.fingerprint
    with CorpusDB(db_path) as db:
        assert db.has_entity(canonical_entity("AquaPump 3000")) == "glossary.yaml"


# ----------------------------------------------------- skipped (explicit)


def test_corrupt_utf8_txt_skipped_rest_ingested(tmp_path):
    folder = tmp_path / "src"
    folder.mkdir()
    (folder / "good.txt").write_text("Returns within 30 days.\n", encoding="utf-8")
    (folder / "bad.txt").write_bytes(b"\xff\xfe garbage \xff bytes")
    result, db_path = _ingest(folder, tmp_path)
    assert result.n_docs == 1
    assert len(result.skipped) == 1
    relpath, reason = result.skipped[0]
    assert relpath == "bad.txt"
    assert "UTF-8" in reason
    with CorpusDB(db_path) as db:
        assert db.contains_text(canonical_text("Returns within 30 days")) == "good.txt"


def test_unsupported_extension_skipped_with_reason(tmp_path):
    folder = tmp_path / "src"
    folder.mkdir()
    (folder / "data.xyz").write_text("whatever", encoding="utf-8")
    (folder / "note.txt").write_text("a note", encoding="utf-8")
    result, _ = _ingest(folder, tmp_path)
    assert result.skipped == (("data.xyz", "unsupported extension .xyz"),)
    assert result.n_docs == 1


def test_hidden_files_and_dirs_silently_ignored(tmp_path):
    folder = tmp_path / "src"
    folder.mkdir()
    (folder / "visible.md").write_text("## Mega Pump\nok\n", encoding="utf-8")
    (folder / ".secret.md").write_text("## Hidden Thing\n", encoding="utf-8")
    hidden_dir = folder / ".git"
    hidden_dir.mkdir()
    (hidden_dir / "notes.md").write_text("## Buried Thing\n", encoding="utf-8")
    result, db_path = _ingest(folder, tmp_path)
    assert result.n_docs == 1
    assert result.skipped == ()  # hidden ≠ skipped-with-reason: never walked
    with CorpusDB(db_path) as db:
        assert db.has_entity(canonical_entity("Hidden Thing")) is None
        assert db.has_entity(canonical_entity("Buried Thing")) is None


def test_db_inside_folder_is_not_ingested_nor_skipped(tmp_path):
    folder = tmp_path / "src"
    folder.mkdir()
    (folder / "note.txt").write_text("Price 12.50 EUR.\n", encoding="utf-8")
    db_path = folder / "corpus.db"
    r1 = ingest_folder(folder, db_path)
    assert r1.n_docs == 1
    assert r1.skipped == ()
    # Reingest with the db file now sitting inside the folder:
    r2 = ingest_folder(folder, db_path)
    assert r2.skipped == ()
    assert r2.fingerprint == r1.fingerprint


def test_subfolder_doc_id_is_posix_relpath(tmp_path):
    folder = tmp_path / "src"
    (folder / "sub").mkdir(parents=True)
    (folder / "sub" / "doc.txt").write_text("Limit is 15%.\n", encoding="utf-8")
    _, db_path = _ingest(folder, tmp_path)
    with CorpusDB(db_path) as db:
        assert db.has_number("percent:15", kind="percent") == "sub/doc.txt"


# ---------------------------------------------------------------- glossary


def test_glossary_yaml_list_form(tmp_path):
    folder = tmp_path / "src"
    folder.mkdir()
    (folder / "note.txt").write_text("nothing here\n", encoding="utf-8")
    (folder / "glossary.yaml").write_text(
        "- AquaPump 3000\n- Méga Widget\n", encoding="utf-8"
    )
    result, db_path = _ingest(folder, tmp_path)
    assert result.n_docs == 2  # note.txt + the glossary registered as a doc
    with CorpusDB(db_path) as db:
        assert db.has_entity(canonical_entity("AquaPump 3000")) == "glossary.yaml"
        assert db.has_entity(canonical_entity("Méga Widget")) == "glossary.yaml"
        assert db.verify_corpus() == (True, [])


def test_glossary_yaml_entities_mapping_form(tmp_path):
    folder = tmp_path / "src"
    folder.mkdir()
    (folder / "note.txt").write_text("nothing here\n", encoding="utf-8")
    (folder / "glossary.yaml").write_text(
        "entities:\n  - HydroFilter Mini\n", encoding="utf-8"
    )
    _, db_path = _ingest(folder, tmp_path)
    with CorpusDB(db_path) as db:
        assert db.has_entity(canonical_entity("HydroFilter Mini")) == "glossary.yaml"


def test_glossary_malformed_yaml_raises_value_error(tmp_path):
    folder = tmp_path / "src"
    folder.mkdir()
    (folder / "glossary.yaml").write_text("entities: [unclosed\n", encoding="utf-8")
    with pytest.raises(ValueError, match="invalid YAML"):
        ingest_folder(folder, tmp_path / "corpus.db")


def test_glossary_wrong_shape_raises_value_error(tmp_path):
    folder = tmp_path / "src"
    folder.mkdir()
    (folder / "glossary.yaml").write_text("just a string\n", encoding="utf-8")
    with pytest.raises(ValueError, match="list of strings"):
        ingest_folder(folder, tmp_path / "corpus.db")
    (folder / "glossary.yaml").write_text("- 42\n", encoding="utf-8")
    with pytest.raises(ValueError, match="non-empty string"):
        ingest_folder(folder, tmp_path / "corpus.db")


def test_glossary_explicit_path_outside_folder(tmp_path):
    folder = tmp_path / "src"
    folder.mkdir()
    (folder / "note.txt").write_text("nothing here\n", encoding="utf-8")
    gloss = tmp_path / "gloss.yaml"
    gloss.write_text("- Turbo Valve X\n", encoding="utf-8")
    result, db_path = _ingest(folder, tmp_path, glossary_path=gloss)
    assert result.n_entities == 1
    with CorpusDB(db_path) as db:
        assert db.has_entity(canonical_entity("Turbo Valve X")) == "glossary:gloss.yaml"
        assert db.verify_corpus() == (True, [])


# -------------------------------------------------------------------- JSON


def test_json_loader_flattens_nested_dict_and_list(tmp_path):
    path = tmp_path / "spec.json"
    path.write_text(
        '{"product": {"sku": "AP-3000-X"}, "tags": ["alpha", "beta"],'
        ' "price": "249.99 EUR"}',
        encoding="utf-8",
    )
    lines = load_file(path).splitlines()
    assert lines == [
        "product.sku: AP-3000-X",
        "tags[0]: alpha",
        "tags[1]: beta",
        "price: 249.99 EUR",
    ]


def test_json_loader_non_dict_roots(tmp_path):
    path = tmp_path / "list.json"
    path.write_text('["AP-3000-X", 42]', encoding="utf-8")
    assert load_file(path) == "[0]: AP-3000-X\n[1]: 42"
    scalar = tmp_path / "scalar.json"
    scalar.write_text('"hello"', encoding="utf-8")
    assert load_file(scalar) == "hello"


def test_json_ingest_extracts_refs_and_numbers_but_no_entities(tmp_path):
    folder = tmp_path / "src"
    folder.mkdir()
    (folder / "catalog.json").write_text(
        '{"name": "AquaPump 3000", "sku": "SKU AP-3000-X", "price": "249.99 EUR"}',
        encoding="utf-8",
    )
    result, db_path = _ingest(folder, tmp_path)
    assert result.n_entities == 0  # JSON is never an entity source
    with CorpusDB(db_path) as db:
        assert db.has_reference(canonical_ref("AP-3000-X")) == "catalog.json"
        assert db.has_number("money:EUR:249.99", kind="money") == "catalog.json"
        assert db.entity_count() == 0


def test_json_corrupt_lands_in_skipped(tmp_path):
    folder = tmp_path / "src"
    folder.mkdir()
    (folder / "broken.json").write_text("{not json", encoding="utf-8")
    result, _ = _ingest(folder, tmp_path)
    assert result.n_docs == 0
    assert len(result.skipped) == 1
    assert result.skipped[0][0] == "broken.json"
    assert "invalid JSON" in result.skipped[0][1]


# --------------------------------------------------------------------- CSV


def test_csv_loader_header_aware_rendering(tmp_path):
    path = tmp_path / "products.csv"
    path.write_text(
        "sku,name,price_eur\nAP-3000-X,AquaPump 3000,249.99\nHF-MINI-2,HydroFilter Mini,39.50\n",
        encoding="utf-8",
    )
    lines = load_file(path).splitlines()
    assert lines[0] == "sku: AP-3000-X | name: AquaPump 3000 | price_eur: 249.99"
    assert lines[1] == "sku: HF-MINI-2 | name: HydroFilter Mini | price_eur: 39.50"


def test_csv_loader_extra_cells_get_positional_labels(tmp_path):
    path = tmp_path / "wide.csv"
    path.write_text("a,b\n1,2,3\n", encoding="utf-8")
    assert load_file(path) == "a: 1 | b: 2 | col3: 3"


def test_csv_name_column_entity_word_count_gate(tmp_path):
    folder = tmp_path / "src"
    folder.mkdir()
    (folder / "items.csv").write_text(
        "id,NAME\n1,Mega Pump\n2,One Two Three Four Five Six Seven\n",
        encoding="utf-8",
    )
    result, db_path = _ingest(folder, tmp_path)
    assert result.n_entities == 1
    with CorpusDB(db_path) as db:
        assert db.has_entity(canonical_entity("Mega Pump")) == "items.csv"
        seven_words = canonical_entity("One Two Three Four Five Six Seven")
        assert db.has_entity(seven_words) is None


def test_csv_money_column_bare_decimals_index_as_money(tmp_path):
    # Real-data eval: ERP exports carry bare decimal prices (no € sign) —
    # without money:EUR indexing, every priced answer was a false positive
    # (56.9% of answers mutilated on such a corpus). D-016.
    folder = tmp_path / "src"
    folder.mkdir()
    (folder / "products.csv").write_text(
        "sku,name,price_eur\nAP-3000-X,AquaPump 3000,249.99\n"
        "HF-MINI-2,HydroFilter Mini,39.50\n",
        encoding="utf-8",
    )
    _, db_path = _ingest(folder, tmp_path)
    with CorpusDB(db_path) as db:
        assert db.has_number("money:EUR:249.99", kind="money") == "products.csv"
        assert db.has_number("money:EUR:39.5", kind="money") == "products.csv"
        # The verbatim rendering still indexes the bare decimal too.
        assert db.has_number("decimal:249.99", kind="decimal") == "products.csv"


def test_csv_money_column_french_multi_comma_amounts(tmp_path):
    # BDPM-style '3,284,71' (thousands commas + decimal comma): quoted in
    # the CSV, parsed cell-wise via canonical_number.
    folder = tmp_path / "src"
    folder.mkdir()
    (folder / "tarifs.csv").write_text(
        'code,libelle,Prix public TTC\nC1,Produit Alpha,"3,284,71"\n'
        'C2,Produit Beta,"12,50"\n',
        encoding="utf-8",
    )
    _, db_path = _ingest(folder, tmp_path)
    with CorpusDB(db_path) as db:
        assert db.has_number("money:EUR:3284.71", kind="money") == "tarifs.csv"
        assert db.has_number("money:EUR:12.5", kind="money") == "tarifs.csv"


def test_csv_money_header_detection_token_vs_substring(tmp_path):
    # 'montant_ht' (substring) and 'total_eur' (whole token) are money;
    # 'couleur' and 'valeur' contain 'eur' only as a substring — NOT money.
    folder = tmp_path / "src"
    folder.mkdir()
    (folder / "data.csv").write_text(
        "ref,montant_ht,total_eur,couleur,valeur\nX1,100.10,200.20,300.30,400.40\n",
        encoding="utf-8",
    )
    _, db_path = _ingest(folder, tmp_path)
    with CorpusDB(db_path) as db:
        assert db.has_number("money:EUR:100.1", kind="money") == "data.csv"
        assert db.has_number("money:EUR:200.2", kind="money") == "data.csv"
        assert db.has_number("money:EUR:300.3", kind="money") is None
        assert db.has_number("money:EUR:400.4", kind="money") is None


def test_csv_money_unparseable_cells_are_skipped_not_indexed(tmp_path):
    # 'N/A', empty, free text, and $-amounts never become money:EUR atoms.
    folder = tmp_path / "src"
    folder.mkdir()
    (folder / "p.csv").write_text(
        "name,price\nAlpha,N/A\nBeta,\nGamma,sur devis\nDelta,$5.00\nEpsilon,8 €\n",
        encoding="utf-8",
    )
    _, db_path = _ingest(folder, tmp_path)
    with CorpusDB(db_path) as db:
        assert db.has_number("money:EUR:8", kind="money") == "p.csv"  # € stripped
        assert db.has_number("money:EUR:5", kind="money") is None  # $ is not EUR
        rows = db._conn.execute(
            "SELECT raw FROM numbers WHERE kind = 'money' AND doc_id = 'p.csv'"
        ).fetchall()
        assert all("N/A" not in r[0] and "devis" not in r[0] for r in rows)


def test_e2e_bare_decimal_price_answer_verifies(tmp_path):
    # End-to-end regression for the 56.9% FP scenario: corpus = ERP-style
    # CSV with bare decimal prices; the answer states the price with a €
    # sign and must verify.
    from verigate.verify.engine import Verifier

    folder = tmp_path / "src"
    folder.mkdir()
    (folder / "products.csv").write_text(
        'name,prix\nProduit Alpha,"3,284,71"\nProduit Beta,12.34\n',
        encoding="utf-8",
    )
    _, db_path = _ingest(folder, tmp_path)
    with CorpusDB(db_path) as db:
        rep = Verifier(db).verify("Le Produit Beta coûte 12,34 € TTC.")
        money = [
            r for r in rep.atoms if r.atom.canonical.startswith("money:EUR:")
        ]
        assert len(money) == 1
        assert money[0].status.value == "verified"
        assert money[0].matched_source == "products.csv"
        rep2 = Verifier(db).verify("Le Produit Alpha coûte 3 284,71 €.")
        money2 = [
            r for r in rep2.atoms if r.atom.canonical.startswith("money:EUR:")
        ]
        assert len(money2) == 1
        assert money2[0].status.value == "verified"


# -------------------------------------------------------------------- DOCX


def test_docx_loader_extracts_paragraphs_and_tables(tmp_path):
    import docx

    path = tmp_path / "spec.docx"
    document = docx.Document()
    document.add_paragraph("Warranty covers 24 months for the AquaPump 3000.")
    table = document.add_table(rows=1, cols=2)
    table.cell(0, 0).text = "SKU"
    table.cell(0, 1).text = "AP-3000-X"
    document.save(str(path))
    text = load_file(path)
    assert "Warranty covers 24 months for the AquaPump 3000." in text
    assert "AP-3000-X" in text


def test_docx_ingest_refs_from_table_cells(tmp_path):
    import docx

    folder = tmp_path / "src"
    folder.mkdir()
    document = docx.Document()
    document.add_paragraph("Product sheet.")
    table = document.add_table(rows=1, cols=1)
    table.cell(0, 0).text = "SKU HF-MINI-2"
    document.save(str(folder / "sheet.docx"))
    _, db_path = _ingest(folder, tmp_path)
    with CorpusDB(db_path) as db:
        assert db.has_reference(canonical_ref("HF-MINI-2")) == "sheet.docx"


# --------------------------------------------------------------------- PDF


def test_pdf_loader_extracts_page_text(tmp_path):
    path = tmp_path / "spec.pdf"
    path.write_bytes(_minimal_pdf("Spec sheet AP-3000-X rated 550 W price 249.99 EUR"))
    text = load_file(path)
    assert "AP-3000-X" in text
    assert "249.99" in text


def test_pdf_corrupt_lands_in_skipped_rest_ingested(tmp_path):
    folder = tmp_path / "src"
    folder.mkdir()
    (folder / "broken.pdf").write_bytes(b"%PDF-1.7\nnot really a pdf at all")
    (folder / "good.txt").write_text("Limit 25%.\n", encoding="utf-8")
    result, db_path = _ingest(folder, tmp_path)
    assert result.n_docs == 1
    assert len(result.skipped) == 1
    relpath, reason = result.skipped[0]
    assert relpath == "broken.pdf"
    assert "PDF" in reason
    with CorpusDB(db_path) as db:
        assert db.has_number("percent:25", kind="percent") == "good.txt"


# ------------------------------------------------------------------ errors


def test_missing_or_non_dir_folder_raises_value_error(tmp_path):
    with pytest.raises(ValueError, match="folder"):
        ingest_folder(tmp_path / "does_not_exist", tmp_path / "corpus.db")
    a_file = tmp_path / "file.txt"
    a_file.write_text("x", encoding="utf-8")
    with pytest.raises(ValueError, match="folder"):
        ingest_folder(a_file, tmp_path / "corpus.db")


def test_load_file_unknown_suffix_raises_loader_error(tmp_path):
    path = tmp_path / "blob.xyz"
    path.write_text("x", encoding="utf-8")
    assert ".xyz" not in SUPPORTED_EXTENSIONS
    with pytest.raises(LoaderError, match="unsupported"):
        load_file(path)


def test_load_file_corrupt_utf8_chains_unicode_error(tmp_path):
    path = tmp_path / "bad.txt"
    path.write_bytes(b"\xff\xfe broken \xff")
    with pytest.raises(LoaderError) as excinfo:
        load_file(path)
    assert isinstance(excinfo.value.__cause__, UnicodeDecodeError)
